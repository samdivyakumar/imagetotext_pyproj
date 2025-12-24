"""
Document processor to reconstruct Word documents with extracted text
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Dict
import logging
import config
from image_extractor import ImageInfo

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process and reconstruct Word documents with OCR text"""
    
    def __init__(self, document: Document, text_placement: str = None):
        """
        Initialize document processor
        
        Args:
            document: python-docx Document object
            text_placement: 'below' to add text below image, 'replace' to replace image with text
        """
        self.document = document
        self.text_placement = text_placement or config.TEXT_PLACEMENT
        
    def add_text_to_document(self, image_texts: Dict[str, str], images: List[ImageInfo]) -> Document:
        """
        Add extracted text to the document
        
        Args:
            image_texts: Dictionary mapping image_id to extracted text
            images: List of ImageInfo objects with position information
            
        Returns:
            Modified Document object
        """
        # Create a mapping of (paragraph_index, run_index) to extracted text
        position_to_text = {}
        for img in images:
            if img.image_id in image_texts:
                key = (img.paragraph_index, img.run_index)
                position_to_text[key] = image_texts[img.image_id]
        
        # Process paragraphs in reverse order to avoid index shifting
        processed_paragraphs = set()
        
        for img in sorted(images, key=lambda x: (x.paragraph_index, x.run_index), reverse=True):
            if img.image_id not in image_texts:
                continue
            
            text = image_texts[img.image_id]
            para_idx = img.paragraph_index
            
            # Skip if no text was extracted
            if not text or text.strip() == "":
                logger.info(f"No text to add for {img.image_id}")
                continue
            
            if self.text_placement == 'below':
                self._add_text_below_image(para_idx, text)
            elif self.text_placement == 'replace':
                self._replace_image_with_text(para_idx, img.run_index, text)
            
            logger.info(f"Added text for {img.image_id} at paragraph {para_idx}")
        
        return self.document
    
    def _add_text_below_image(self, para_idx: int, text: str):
        """Add text in a new paragraph below the image"""
        try:
            # Get the paragraph containing the image
            image_paragraph = self.document.paragraphs[para_idx]
            
            # Insert a new paragraph after the image paragraph
            # We need to work with the underlying XML element
            new_para = self.document.add_paragraph()
            
            # Move the new paragraph to the correct position
            # Get the parent element
            parent = image_paragraph._element.getparent()
            # Insert after the image paragraph
            parent.insert(parent.index(image_paragraph._element) + 1, new_para._element)
            
            # Add the prefix
            if config.TEXT_PREFIX:
                prefix_run = new_para.add_run(config.TEXT_PREFIX)
                prefix_run.bold = True
                prefix_run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
            
            # Add the extracted text
            text_run = new_para.add_run(text)
            text_run.font.size = Pt(10)
            
            # Add the suffix
            if config.TEXT_SUFFIX:
                suffix_run = new_para.add_run(config.TEXT_SUFFIX)
                suffix_run.bold = True
                suffix_run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
            
        except Exception as e:
            logger.error(f"Failed to add text below image at paragraph {para_idx}: {e}")
    
    def _replace_image_with_text(self, para_idx: int, run_idx: int, text: str):
        """Replace the image with extracted text"""
        try:
            # Get the paragraph containing the image
            paragraph = self.document.paragraphs[para_idx]
            
            # Clear the run containing the image
            if run_idx < len(paragraph.runs):
                run = paragraph.runs[run_idx]
                run.clear()
                
                # Add the prefix
                if config.TEXT_PREFIX:
                    prefix_run = paragraph.add_run(config.TEXT_PREFIX)
                    prefix_run.bold = True
                    prefix_run.font.color.rgb = RGBColor(0, 100, 0)
                
                # Add the text
                text_run = paragraph.add_run(text)
                text_run.font.size = Pt(10)
                
                # Add the suffix
                if config.TEXT_SUFFIX:
                    suffix_run = paragraph.add_run(config.TEXT_SUFFIX)
                    suffix_run.bold = True
                    suffix_run.font.color.rgb = RGBColor(0, 100, 0)
            
        except Exception as e:
            logger.error(f"Failed to replace image at paragraph {para_idx}, run {run_idx}: {e}")
    
    def save_document(self, output_path: str):
        """Save the modified document"""
        try:
            self.document.save(output_path)
            logger.info(f"Document saved successfully to: {output_path}")
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            raise
